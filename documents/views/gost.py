# documents/views.py (или documents/views/gost.py, если вы разделяете)

from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from django.views.generic import ListView, DetailView, View, UpdateView # Добавлены View и UpdateView
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin # Используем миксины для CBV
from django.contrib.auth.decorators import login_required # Оставим для примера экспорта (если он будет)
from django.http import HttpResponse # Для экспорта
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
import re

# Импортируем модели и формы (пути могут отличаться в зависимости от структуры вашего проекта)
# Убедитесь, что эти импорты верны для вашей структуры:
# from .models.gost import Document, Abstract, Performer, Term, Abbreviation, Reference, Appendix # Пример, если модели в models/gost.py
# from .forms.gost import DocumentForm, TitlePageForm, AbstractForm, PerformerFormSet, TermFormSet, AbbreviationFormSet, ReferenceFormSet, AppendixFormSet # Пример, если формы в forms/gost.py
# Если все в одном файле models.py и forms.py:
from documents.models.gost import Document,Abstract
from documents.forms.gost import DocumentForm,AppendixFormSet,TermFormSet,AbbreviationFormSet,PerformerFormSet,ReferenceFormSet,TitlePageForm,AbstractForm


# --- Классовые представления (CBV) ---

class DocumentListView(LoginRequiredMixin, ListView):
    """
    Отображает список документов, принадлежащих текущему пользователю.
    """
    model = Document
    template_name = 'documents/document_list.html'
    context_object_name = 'documents' # Имя переменной в шаблоне

    def get_queryset(self):
        """
        Возвращает только документы, принадлежащие текущему пользователю.
        """
        # Поле в модели Document называется 'user', а не 'owner'
        return Document.objects.filter(user=self.request.user).order_by('-created_at')

class DocumentDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Document
    template_name = 'documents/document_detail.html'
    context_object_name = 'document'

    def test_func(self):
        # Проверяем, что пользователь является владельцем документа
        document = self.get_object()
        return document.user == self.request.user

class DocumentCreateView(LoginRequiredMixin, View):
    """
    Обрабатывает создание нового документа со всеми связанными формами и формсетами.
    Используем базовый View для гибкости.
    """
    template_name = 'documents/document_form.html'

    def get(self, request, *args, **kwargs):
        """Обработка GET-запроса: отображение пустых форм."""
        doc_form = DocumentForm()
        # Важно: Убедитесь, что TitlePageForm и AbstractForm не требуют instance документа при инициализации
        # Если требуют, их нужно будет инициализировать в POST после создания document
        title_form = TitlePageForm()
        abstract_form = AbstractForm()
        performer_formset = PerformerFormSet(prefix='performers')
        term_formset = TermFormSet(prefix='terms')
        abbrev_formset = AbbreviationFormSet(prefix='abbrevs')
        reference_formset = ReferenceFormSet(prefix='refs')
        appendix_formset = AppendixFormSet(prefix='apps')

        context = {
            'form': doc_form,
            'title_form': title_form,
            'abstract_form': abstract_form,
            'performer_formset': performer_formset,
            'term_formset': term_formset,
            'abbrev_formset': abbrev_formset,
            'reference_formset': reference_formset,
            'appendix_formset': appendix_formset,
        }
        return render(request, self.template_name, context)

    def post(self, request, *args, **kwargs):
        """Обработка POST-запроса: валидация и сохранение данных."""
        print("Данные POST:", request.POST)
        doc_form = DocumentForm(request.POST)
        # Для связанных форм (OneToOne) обычно передают request.POST без instance
        title_form = TitlePageForm(request.POST)
        abstract_form = AbstractForm(request.POST)
        # Для формсетов передаем request.POST и префикс
        performer_formset = PerformerFormSet(request.POST, request.FILES, prefix='performers') # Добавил request.FILES на всякий случай
        term_formset = TermFormSet(request.POST, request.FILES, prefix='terms')
        abbrev_formset = AbbreviationFormSet(request.POST, request.FILES, prefix='abbrevs')
        reference_formset = ReferenceFormSet(request.POST, request.FILES, prefix='refs')
        appendix_formset = AppendixFormSet(request.POST, request.FILES, prefix='apps')

        # Проверка валидности форм
        forms_valid = all([
            doc_form.is_valid(),
            title_form.is_valid(),
            abstract_form.is_valid(),
            performer_formset.is_valid(),
            term_formset.is_valid(),
            abbrev_formset.is_valid(),
            reference_formset.is_valid(),
            appendix_formset.is_valid(),
        ])
        
        print(f"Формы валидны: {forms_valid}")
        
        # Выводим информацию об ошибках для отладки
        if not doc_form.is_valid():
            print(f"Ошибки DocumentForm: {doc_form.errors}")
        if not title_form.is_valid():
            print(f"Ошибки TitlePageForm: {title_form.errors}")
        if not abstract_form.is_valid():
            print(f"Ошибки AbstractForm: {abstract_form.errors}")
        
        if forms_valid:
            try:
                # Сохраняем основной документ
                document = doc_form.save(commit=False)
                document.user = request.user # Устанавливаем владельца
                
                # Добавляем дополнительную проверку
                if not request.user.is_authenticated:
                    raise Exception("Пользователь не аутентифицирован. Пожалуйста, войдите в систему.")
                
                document.save() # Сохраняем, чтобы получить PK
                
                # Выводим отладочную информацию
                print(f"Документ сохранен: {document.id} - {document.title} (владелец: {document.user.username})")

                # Сохраняем связанные объекты OneToOne
                title_page = title_form.save(commit=False)
                title_page.document = document
                title_page.save()
                print(f"Титульный лист сохранен: {title_page.id}")

                abstract = abstract_form.save(commit=False)
                abstract.document = document
                abstract.save()
                print(f"Реферат сохранен: {abstract.id}")

                # Сохраняем формсеты (связанные объекты ForeignKey)
                formsets_to_save = [
                    performer_formset, term_formset, abbrev_formset,
                    reference_formset, appendix_formset
                ]
                for formset in formsets_to_save:
                    # Сохраняем формсет, он автоматически обрабатывает удаления
                    instances = formset.save(commit=False)
                    for instance in instances:
                        instance.document = document
                        instance.save()
                    # Сохраняем возможные M2M связи
                    formset.save_m2m()

                # Перенаправляем на страницу просмотра
                return redirect('document_detail', pk=document.pk)
            except Exception as e:
                print(f"Ошибка при сохранении документа: {str(e)}")
                import traceback
                traceback.print_exc()
                # Возвращаем ошибку для отображения пользователю
                context = {
                    'form': doc_form,
                    'title_form': title_form,
                    'abstract_form': abstract_form,
                    'performer_formset': performer_formset,
                    'term_formset': term_formset,
                    'abbrev_formset': abbrev_formset,
                    'reference_formset': reference_formset,
                    'appendix_formset': appendix_formset,
                    'error_message': f"Ошибка при сохранении: {str(e)}"
                }
                return render(request, self.template_name, context)
        else:
            # Выводим ошибки валидации форм
            print("Ошибки валидации форм:")
            if not doc_form.is_valid():
                print(f"Ошибки doc_form: {doc_form.errors}")
            if not title_form.is_valid():
                print(f"Ошибки title_form: {title_form.errors}")
            if not abstract_form.is_valid():
                print(f"Ошибки abstract_form: {abstract_form.errors}")
            if not performer_formset.is_valid():
                print(f"Ошибки performer_formset: {performer_formset.errors}")
            if not term_formset.is_valid():
                print(f"Ошибки term_formset: {term_formset.errors}")
            if not abbrev_formset.is_valid():
                print(f"Ошибки abbrev_formset: {abbrev_formset.errors}")
            if not reference_formset.is_valid():
                print(f"Ошибки reference_formset: {reference_formset.errors}")
            if not appendix_formset.is_valid():
                print(f"Ошибки appendix_formset: {appendix_formset.errors}")
                
            # Если хотя бы одна форма/формсет невалидны, рендерим страницу снова с ошибками
            context = {
                'form': doc_form,
                'title_form': title_form,
                'abstract_form': abstract_form,
                'performer_formset': performer_formset,
                'term_formset': term_formset,
                'abbrev_formset': abbrev_formset,
                'reference_formset': reference_formset,
                'appendix_formset': appendix_formset,
                'error_message': "Пожалуйста, исправьте ошибки в форме."
            }
            return render(request, self.template_name, context)


class DocumentUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    """
    Обрабатывает редактирование документа и всех связанных объектов.
    """
    model = Document
    form_class = DocumentForm
    template_name = 'documents/document_form.html'
    pk_url_kwarg = 'pk'
    context_object_name = 'document'

    def test_func(self):
        """Проверяет, что текущий пользователь является владельцем документа."""
        document = self.get_object()
        return document.user == self.request.user

    def get_context_data(self, **kwargs):
        """
        Добавляем связанные формы и формсеты в контекст с данными
        """
        context = super().get_context_data(**kwargs)
        
        if self.request.POST:
            # Если POST-запрос, инициализируем формы с данными запроса
            if 'title_form' not in context:
                context['title_form'] = TitlePageForm(
                    self.request.POST,
                    instance=getattr(self.object, 'title_page', None)
                )
            if 'abstract_form' not in context:
                context['abstract_form'] = AbstractForm(
                    self.request.POST,
                    instance=getattr(self.object, 'abstract', None)
                )
            
            # Инициализируем формсеты с данными POST
            if 'performer_formset' not in context:
                context['performer_formset'] = PerformerFormSet(
                    self.request.POST, 
                    self.request.FILES,
                    prefix='performers',
                    queryset=self.object.performers.all()
                )
            if 'term_formset' not in context:
                context['term_formset'] = TermFormSet(
                    self.request.POST, 
                    self.request.FILES,
                    prefix='terms',
                    queryset=self.object.terms.all()
                )
            if 'abbrev_formset' not in context:
                context['abbrev_formset'] = AbbreviationFormSet(
                    self.request.POST, 
                    self.request.FILES,
                    prefix='abbrevs',
                    queryset=self.object.abbreviations.all()
                )
            if 'reference_formset' not in context:
                context['reference_formset'] = ReferenceFormSet(
                    self.request.POST, 
                    self.request.FILES,
                    prefix='refs',
                    queryset=self.object.references.all()
                )
            if 'appendix_formset' not in context:
                context['appendix_formset'] = AppendixFormSet(
                    self.request.POST, 
                    self.request.FILES,
                    prefix='apps',
                    queryset=self.object.appendices.all()
                )
        else:
            # Если GET-запрос, инициализируем формы с данными из БД
            if 'title_form' not in context:
                try:
                    context['title_form'] = TitlePageForm(instance=self.object.title_page)
                except (AttributeError, Exception):
                    context['title_form'] = TitlePageForm()
                
            if 'abstract_form' not in context:
                try:
                    context['abstract_form'] = AbstractForm(instance=self.object.abstract)
                except (Abstract.DoesNotExist, Exception):
                    context['abstract_form'] = AbstractForm()
            
            # Инициализируем формсеты с данными из БД
            if 'performer_formset' not in context:
                context['performer_formset'] = PerformerFormSet(
                    prefix='performers',
                    queryset=self.object.performers.all()
                )
            if 'term_formset' not in context:
                context['term_formset'] = TermFormSet(
                    prefix='terms',
                    queryset=self.object.terms.all()
                )
            if 'abbrev_formset' not in context:
                context['abbrev_formset'] = AbbreviationFormSet(
                    prefix='abbrevs',
                    queryset=self.object.abbreviations.all()
                )
            if 'reference_formset' not in context:
                context['reference_formset'] = ReferenceFormSet(
                    prefix='refs',
                    queryset=self.object.references.all()
                )
            if 'appendix_formset' not in context:
                context['appendix_formset'] = AppendixFormSet(
                    prefix='apps',
                    queryset=self.object.appendices.all()
                )

        return context

    def form_valid(self, form):
        """
        Обрабатывает валидные данные формы и всех связанных форм/формсетов.
        """
        context = self.get_context_data()
        title_form = context['title_form']
        abstract_form = context['abstract_form']
        performer_formset = context['performer_formset']
        term_formset = context['term_formset']
        abbrev_formset = context['abbrev_formset']
        reference_formset = context['reference_formset']
        appendix_formset = context['appendix_formset']

        # Проверяем валидность всех форм и формсетов
        if all([
            form.is_valid(),
            title_form.is_valid(),
            abstract_form.is_valid(),
            performer_formset.is_valid(),
            term_formset.is_valid(),
            abbrev_formset.is_valid(),
            reference_formset.is_valid(),
            appendix_formset.is_valid(),
        ]):
            # Сохраняем основной документ
            self.object = form.save()

            # Сохраняем связанные объекты OneToOne
            # Для title_page
            title_page = title_form.save(commit=False)
            title_page.document = self.object
            title_page.save()

            # Для abstract
            abstract = abstract_form.save(commit=False)
            abstract.document = self.object
            abstract.save()

            # Сохраняем формсеты (связанные объекты ForeignKey)
            formsets_to_save = [
                performer_formset, term_formset, abbrev_formset,
                reference_formset, appendix_formset
            ]
            for formset in formsets_to_save:
                # Сохраняем формсет, он автоматически обрабатывает удаления
                instances = formset.save(commit=False)
                for instance in instances:
                    instance.document = self.object
                    instance.save()
                # Сохраняем возможные M2M связи
                formset.save_m2m()

            return super().form_valid(form)
        else:
            return self.form_invalid(form)

    def get_success_url(self):
        """Возвращает URL для перенаправления после успешного обновления."""
        return reverse_lazy('document_detail', kwargs={'pk': self.object.pk})

# --- Функции для экспорта ---
@login_required
def document_export_docx(request, pk):
    """
    Экспортирует документ в формат DOCX
    """
    try:
        # Получаем документ из базы данных
        document = get_object_or_404(Document, pk=pk, user=request.user)
        
        # Создаем docx документ
        docx = DocxDocument()
        
        # Настройка стилей документа
        styles = docx.styles
        
        # Стиль для заголовков
        style_heading1 = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        style_heading1.base_style = styles['Heading 1']
        font = style_heading1.font
        font.name = 'Times New Roman'
        font.size = Pt(16)
        font.bold = True
        
        # Стиль для текста
        style_normal = styles['Normal']
        font = style_normal.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        # Добавляем титульный лист
        add_title_page(docx, document)

        # Добавляем реферат
        add_abstract(docx, document)
        
        # Добавляем основные разделы
        add_sections(docx, document)
        
        # Сохраняем файл
        buffer = BytesIO()
        docx.save(buffer)
        buffer.seek(0)
        
        # Отправляем файл пользователю
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        file_name = f"{document.title.replace(' ', '_')}.docx"
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        
        return response
        
    except Exception as e:
        print(f"Ошибка при экспорте документа: {str(e)}")
        return HttpResponse(f"Ошибка при экспорте документа: {str(e)}", status=500)

def add_title_page(docx, document):
    """
    Добавляет титульный лист в docx документ
    """
    # Получаем объект титульного листа
    try:
        title_page = document.title_page
    except:
        title_page = None
    
    # Добавляем информацию титульного листа
    p = docx.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_page and title_page.department:
        p.add_run(title_page.department.upper()).bold = True
    
    # Добавляем отступ
    for _ in range(3):
        docx.add_paragraph()
    
    # Заголовок документа
    p = docx.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(document.title.upper()).bold = True
    
    # Тип отчета
    if document.report_type:
        p = docx.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_type_display = dict(document._meta.get_field('report_type').choices)[document.report_type]
        p.add_run(f"{report_type_display.upper()} ОТЧЁТ").bold = True
    
    # Добавляем информацию о руководителе и исполнителях в нижней части титульного листа
    for _ in range(5):
        docx.add_paragraph()
    
    # Информация о руководителе
    if title_page:
        table = docx.add_table(rows=2, cols=2)
        table.autofit = True
        
        # Левая сторона - руководитель
        if title_page.head_position or title_page.head_full_name:
            cell = table.cell(0, 0)
            cell.text = title_page.head_position or "Руководитель проекта"
            
            cell = table.cell(1, 0)
            cell.text = title_page.head_full_name or ""
    
    # Добавляем город и год в нижней части страницы
    for _ in range(3):
        docx.add_paragraph()
    
    p = docx.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{document.year}").bold = True
    
    # Добавляем разрыв страницы
    docx.add_page_break()

def add_abstract(docx, document):
    """
    Добавляет реферат в docx документ
    """
    # Заголовок
    p = docx.add_paragraph("РЕФЕРАТ", style='CustomHeading1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Получаем объект реферата
    try:
        abstract = document.abstract
        
        # Добавляем содержимое реферата
        if abstract and abstract.content:
            # Очищаем HTML теги
            clean_text = re.sub(r'<.*?>', '', abstract.content)
            p = docx.add_paragraph(clean_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except:
        # Если реферат не найден, добавляем пустой параграф
        docx.add_paragraph()
    
    # Разрыв страницы
    docx.add_page_break()

def add_sections(docx, document):
    """
    Добавляет основные разделы документа
    """
    # Содержание
    p = docx.add_paragraph("СОДЕРЖАНИЕ", style='CustomHeading1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # TODO: Автоматически генерировать содержание
    docx.add_paragraph()
    docx.add_page_break()
    
    # Введение
    p = docx.add_paragraph("ВВЕДЕНИЕ", style='CustomHeading1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if document.introduction:
        # Очищаем HTML теги
        clean_text = re.sub(r'<.*?>', '', document.introduction)
        p = docx.add_paragraph(clean_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        docx.add_paragraph()
    
    docx.add_page_break()
    
    # Основная часть
    p = docx.add_paragraph("ОСНОВНАЯ ЧАСТЬ", style='CustomHeading1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if document.main_part:
        # Очищаем HTML теги
        clean_text = re.sub(r'<.*?>', '', document.main_part)
        p = docx.add_paragraph(clean_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        docx.add_paragraph()
    
    docx.add_page_break()
    
    # Заключение
    p = docx.add_paragraph("ЗАКЛЮЧЕНИЕ", style='CustomHeading1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if document.conclusion:
        # Очищаем HTML теги
        clean_text = re.sub(r'<.*?>', '', document.conclusion)
        p = docx.add_paragraph(clean_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        docx.add_paragraph()
    
    docx.add_page_break()
    
    # Список использованных источников
    p = docx.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", style='CustomHeading1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем список источников
    try:
        references = document.references.all().order_by('order')
        if references.exists():
            for i, ref in enumerate(references, 1):
                p = docx.add_paragraph(f"{i}. {ref.citation}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            docx.add_paragraph()
    except:
        docx.add_paragraph()
    
    # Добавляем приложения, если они есть
    try:
        appendices = document.appendices.all().order_by('order')
        if appendices.exists():
            docx.add_page_break()
            
            for appendix in appendices:
                p = docx.add_paragraph(f"ПРИЛОЖЕНИЕ {appendix.label}", style='CustomHeading1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if appendix.title:
                    p = docx.add_paragraph(appendix.title)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style = 'CustomHeading1'
                
                if appendix.content:
                    # Очищаем HTML теги
                    clean_text = re.sub(r'<.*?>', '', appendix.content)
                    p = docx.add_paragraph(clean_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                docx.add_page_break()
    except:
        pass