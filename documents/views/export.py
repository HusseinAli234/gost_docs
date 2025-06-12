from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect
from django.utils.text import slugify
from django.conf import settings
from django.contrib import messages
import io
from io import BytesIO, StringIO
import os
import logging
import tempfile
import base64
import re
import mimetypes
import json
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert
# Импорт DocxTemplate для работы с шаблонами
from docxtpl import DocxTemplate
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.image import ImagePart

# Импортируем htmldocx вместо mammoth
from htmldocx import HtmlToDocx

# Попытка импортировать pythoncom для Windows
try:
    import pythoncom
except ImportError:
    pythoncom = None # Устанавливаем в None, если импорт не удался (не Windows)

from documents.models.sto import Document_sto
from documents.models.main import Document_main

# Импортируем модуль AI для получения стилей форматирования
import ai

# Настройка логгера
logger = logging.getLogger(__name__)

# Отображение типов работ на названия шаблонов
WORK_TYPE_TEMPLATES = {
    'MAG_DIPLOMA': 'magitr_dissertation',
    'DIPLOMA': 'diplom_work',
    'BACHELOR': 'bachelor_work',
    'COURSE': 'kursovaya',
    'CALC_GRAPH': 'diplo_project',  # используем общий шаблон
    'PRACTICE': 'otchet_praktika',
    'LAB': 'laba_work',
    'REF': 'referat',
}

# Импорты для низкоуровневой работы с OXML элементами
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

def clean_html(html_content):
    """Очищает HTML от тегов и возвращает только текст"""
    if not html_content:
        return ""
    return BeautifulSoup(html_content, 'html.parser').get_text()

def django_mammoth_image_converter(image):
    """
    Пользовательская функция-конвертер изображений для mammoth.
    Mammoth вызывает ее для каждого тега <img>.
    Получает объект mammoth.images.Image, представляющий тег.
    Должна вернуть словарь с ключами 'data' (байты изображения) и 'contentType' (MIME-тип).
    Обрабатывает пути к медиа-файлам Django.
    """
    src = image.src
    logger.debug(f"Mammoth встретил изображение с src: {src}")

    try:
        # Проверяем различные типы источников изображений
        if src and src.startswith(settings.MEDIA_URL):
            # 1. Изображение из медиа-хранилища Django
            relative_media_path = src[len(settings.MEDIA_URL):]
            file_path = os.path.join(settings.MEDIA_ROOT, relative_media_path)

            if os.path.exists(file_path):
                with open(file_path, 'rb') as f:
                    image_data = f.read()
                content_type, _ = mimetypes.guess_type(file_path)
                logger.debug(f"Загружено изображение из медиа: {file_path}, тип: {content_type}")
            else:
                logger.warning(f"Файл изображения не найден: {file_path}")
                return None

        elif src and src.startswith('data:image/'):
            # 2. Изображение в формате base64
            try:
                header, encoded = src.split(",", 1)
                content_type = header.split(";")[0].split(":")[1]
                image_data = base64.b64decode(encoded)
                logger.debug(f"Декодировано base64 изображение, тип: {content_type}")
            except Exception as e:
                logger.error(f"Ошибка при декодировании base64 изображения: {e}")
                return None

        elif src and (src.startswith('http://') or src.startswith('https://')):
            # 3. Изображение по внешнему URL
            try:
                response = requests.get(src, stream=True, timeout=10)
                if response.status_code == 200:
                    image_data = response.content
                    content_type = response.headers.get('Content-Type', 'image/jpeg')
                    logger.debug(f"Загружено изображение по URL: {src}, тип: {content_type}")
                else:
                    logger.warning(f"Не удалось загрузить изображение по URL: {src}, статус: {response.status_code}")
                    return None
            except Exception as e:
                logger.error(f"Ошибка при загрузке изображения по URL {src}: {e}")
                return None
        else:
            logger.warning(f"Неподдерживаемый формат src изображения: {src}")
            return None

        # Проверяем, что мы определили MIME-тип
        if content_type is None:
            logger.warning(f"Не удалось определить MIME-тип изображения: {src}. Изображение будет пропущено.")
            return None

        # Возвращаем данные в формате, который ожидает mammoth
        return {
            'src': src,
            'alt': image.alt,
            'data': image_data,
            'contentType': content_type
        }

    except Exception as e:
        logger.error(f"Общая ошибка при обработке изображения {src}: {e}", exc_info=True)
        return None

def process_html_to_docx(html_content, docx_document):
    """
    Преобразует HTML-контент в DOCX и добавляет его в существующий документ.
    Изображения загружаются напрямую по URL и вставляются в правильные места,
    сохраняя текст вокруг изображений.
    
    Args:
        html_content (str): HTML-контент для преобразования.
        docx_document (Document): Существующий объект python-docx, куда будет добавлен контент.
    """
    if not html_content:
        logger.info("HTML контент пуст, нечего добавлять.")
        return

    logger.info(f"Обработка HTML контента длиной {len(html_content)} символов")
    
    # Сохраняем информацию об отступах из HTML перед парсингом
    # Ищем теги <p> с отступами через регулярные выражения
    indent_info = {}
    p_tags = re.finditer(r'<p\s+([^>]*)style="([^"]*)"', html_content)
    for i, match in enumerate(p_tags):
        attrs = match.group(1)
        style = match.group(2)
        
        # Генерируем уникальный ID для параграфа
        p_id = f"p_{i}"
        
        # Извлекаем значения отступов из стилей
        text_indent = re.search(r'text-indent:\s*([0-9.]+)(px|em|cm|mm|pt)', style)
        margin_left = re.search(r'margin-left:\s*([0-9.]+)(px|em|cm|mm|pt)', style)
        
        if text_indent or margin_left:
            indent_info[p_id] = {
                'has_indent': True,
                'text_indent': text_indent.groups() if text_indent else None,
                'margin_left': margin_left.groups() if margin_left else None
            }
            
            # Добавляем ID к тегу <p> для последующей идентификации
            html_content = html_content.replace(match.group(0), f'<p id="{p_id}" {attrs}style="{style}"')
    
    # Дополнительно ищем параграфы с классом has-indent
    p_with_class = re.finditer(r'<p\s+([^>]*)class="([^"]*has-indent[^"]*)"([^>]*)>', html_content)
    for i, match in enumerate(p_with_class, start=len(indent_info)):
        before_class = match.group(1)
        class_attr = match.group(2)
        after_class = match.group(3)
        
        # Генерируем уникальный ID для параграфа
        p_id = f"p_indent_class_{i}"
        
        # Добавляем информацию об отступе (стандартный отступ 1.25 см)
        indent_info[p_id] = {
            'has_indent': True,
            'text_indent': ('1.25', 'cm'),
            'margin_left': None
        }
        
        # Добавляем ID к тегу <p> для последующей идентификации
        html_content = html_content.replace(
            match.group(0), 
            f'<p id="{p_id}" {before_class}class="{class_attr}"{after_class}>'
        )
    
    logger.info(f"Найдено {len(indent_info)} параграфов с отступами")
    
    # Используем BeautifulSoup для парсинга HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Проверяем наличие неразрывных пробелов в тексте (уже преобразованных BeautifulSoup)
    paragraphs_with_nbsp = {}
    for p in soup.find_all('p'):
        p_id = p.get('id', f"auto_{id(p)}")
        
        # Проверяем текст на наличие неразрывных пробелов в начале
        text = p.get_text()
        nbsp_count = 0
        for char in text:
            if char == '\u00A0':  # Неразрывный пробел
                nbsp_count += 1
            else:
                break
                
        if nbsp_count > 0:
            paragraphs_with_nbsp[p_id] = nbsp_count
            logger.info(f"Параграф {p_id} содержит {nbsp_count} неразрывных пробелов в начале")
    
    # Преобразуем обратно в строку для дальнейшей обработки
    html_content = str(soup)
    
    # Патч для библиотеки htmldocx
    try:
        from htmldocx.h2d import HtmlToDocx as OriginalHtmlToDocx
        
        # Создаем патч класса HtmlToDocx
        class PatchedHtmlToDocx(OriginalHtmlToDocx):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self.run = None  # Добавляем отсутствующий атрибут
                self.current_paragraph_id = None  # Для отслеживания ID текущего параграфа
                self.paragraphs_created = []  # Список созданных параграфов
            
            def handle_starttag(self, tag, attrs):
                try:
                    # Сохраняем ID для параграфов
                    if tag == 'p':
                        self.current_paragraph_id = None
                        for attr in attrs:
                            if attr[0] == 'id':
                                self.current_paragraph_id = attr[1]
                                break
                        
                        # Создаем новый параграф
                        old_paragraph = self.paragraph
                        super().handle_starttag(tag, attrs)
                        
                        # Если был создан новый параграф, сохраняем его и его ID
                        if self.paragraph != old_paragraph and self.paragraph is not None:
                            self.paragraphs_created.append((self.paragraph, self.current_paragraph_id))
                    else:
                        # Пытаемся выполнить оригинальный метод
                        super().handle_starttag(tag, attrs)
                        
                    # Проверяем, нужно ли инициализировать self.run
                    if tag in ['br', 'hr'] and self.run is None and self.paragraph is not None:
                        # Создаем run, если его нет
                        self.run = self.paragraph.add_run()
                        
                except AttributeError as e:
                    # Если возникла ошибка с атрибутом run, игнорируем ее
                    if "'NoneType' object has no attribute 'add_break'" in str(e):
                        logger.warning("Игнорирована ошибка в htmldocx: 'NoneType' object has no attribute 'add_break'")
                        # Пытаемся создать run и повторить операцию
                        if self.paragraph is not None:
                            self.run = self.paragraph.add_run()
                            if tag == 'br':
                                self.run.add_break()
                    elif "'NoneType' object has no attribute" in str(e):
                        logger.warning(f"Игнорирована ошибка в htmldocx: {str(e)}")
                    else:
                        # Другие ошибки пробрасываем дальше
                        raise
        
        # Заменяем оригинальный класс на патченный
        HtmlToDocx = PatchedHtmlToDocx
        logger.info("Применен расширенный патч для библиотеки htmldocx")
    except Exception as e:
        logger.warning(f"Не удалось применить патч для htmldocx: {e}")
    
    try:
        # Создаем временную директорию для хранения изображений
        with tempfile.TemporaryDirectory() as temp_dir:
            # Подготавливаем изображения и сохраняем их во временную директорию
            soup = BeautifulSoup(html_content, 'html.parser')
            images_map = {}
            
            # Находим все изображения
            images = soup.find_all('img')
            for i, img in enumerate(images):
                src = img.get('src', '')
                if not src:
                    continue
                
                logger.info(f"Обрабатываю изображение {i+1}/{len(images)}: {src[:50]}...")
                
                # Генерируем уникальный идентификатор для этого изображения
                img_id = f"IMG_PLACEHOLDER_{i}"
                
                # Пытаемся загрузить изображение
                img_path = None
                
                try:
                    # Обрабатываем различные форматы src
                    if src.startswith(('http://', 'https://')):
                        # Внешний URL
                        logger.info(f"Загрузка изображения по URL: {src}")
                        response = requests.get(src, stream=True, timeout=10)
                        if response.status_code == 200:
                            # Определяем расширение файла из Content-Type или URL
                            content_type = response.headers.get('Content-Type', '')
                            ext = mimetypes.guess_extension(content_type) or '.png'
                            if ext == '.jpe':
                                ext = '.jpg'
                            
                            # Сохраняем изображение во временную директорию
                            img_path = os.path.join(temp_dir, f"image_{i}{ext}")
                            with open(img_path, 'wb') as f:
                                f.write(response.content)
                            logger.info(f"Изображение сохранено: {img_path}")
                        else:
                            logger.warning(f"Не удалось загрузить изображение, статус: {response.status_code}")
                    
                    elif src.startswith('data:image/'):
                        # Data URL (base64)
                        logger.info("Обработка Data URL изображения")
                        try:
                            header, encoded = src.split(",", 1)
                            content_type = header.split(";")[0].split(":")[1]
                            ext = mimetypes.guess_extension(content_type) or '.png'
                            if ext == '.jpe':
                                ext = '.jpg'
                            
                            # Декодируем base64 и сохраняем изображение
                            img_path = os.path.join(temp_dir, f"image_{i}{ext}")
                            with open(img_path, 'wb') as f:
                                f.write(base64.b64decode(encoded))
                            logger.info(f"Base64 изображение сохранено: {img_path}")
                        except Exception as e:
                            logger.error(f"Ошибка при декодировании base64: {e}")
                    
                    elif src.startswith('/'):
                        # Локальный путь от корня сайта
                        logger.info(f"Обработка локального пути: {src}")
                        
                        # Если путь начинается с /media/, используем MEDIA_ROOT
                        if src.startswith('/media/'):
                            file_path = os.path.join(settings.MEDIA_ROOT, src[7:])
                        # Иначе пробуем относительно BASE_DIR
                        else:
                            file_path = os.path.join(settings.BASE_DIR, src[1:])
                        
                        if os.path.exists(file_path):
                            # Копируем файл во временную директорию
                            ext = os.path.splitext(file_path)[1] or '.png'
                            img_path = os.path.join(temp_dir, f"image_{i}{ext}")
                            with open(file_path, 'rb') as src_file, open(img_path, 'wb') as dst_file:
                                dst_file.write(src_file.read())
                            logger.info(f"Локальное изображение скопировано: {img_path}")
                        else:
                            # Если файл не найден, пробуем построить URL и загрузить через HTTP
                            try:
                                full_url = f"http://localhost:8000{src}"
                                logger.info(f"Пробую загрузить через HTTP: {full_url}")
                                response = requests.get(full_url, stream=True, timeout=10)
                                if response.status_code == 200:
                                    content_type = response.headers.get('Content-Type', '')
                                    ext = mimetypes.guess_extension(content_type) or '.png'
                                    if ext == '.jpe':
                                        ext = '.jpg'
                                    
                                    img_path = os.path.join(temp_dir, f"image_{i}{ext}")
                                    with open(img_path, 'wb') as f:
                                        f.write(response.content)
                                    logger.info(f"Изображение загружено через HTTP: {img_path}")
                                else:
                                    logger.warning(f"Не удалось загрузить через HTTP, статус: {response.status_code}")
                            except Exception as e:
                                logger.error(f"Ошибка при HTTP загрузке: {e}")
                    
                    # Если удалось загрузить изображение, сохраняем информацию
                    if img_path and os.path.exists(img_path):
                        images_map[img_id] = img_path
                        # Заменяем тег img на специальный маркер, который не будет интерпретирован как HTML
                        img_marker = soup.new_string(f"[[IMG:{img_id}]]")
                        img.replace_with(img_marker)
                
                except Exception as e:
                    logger.error(f"Ошибка при обработке изображения {src}: {e}", exc_info=True)
            
            # Преобразуем модифицированный HTML в строку
            modified_html = str(soup)
            
            # Разбиваем HTML на части по маркерам изображений
            parts = []
            last_pos = 0
            
            for match in re.finditer(r'\[\[IMG:(IMG_PLACEHOLDER_\d+)\]\]', modified_html):
                img_id = match.group(1)
                start_pos = match.start()
                end_pos = match.end()
                
                # Добавляем текст до изображения
                if start_pos > last_pos:
                    parts.append(('text', modified_html[last_pos:start_pos]))
                
                # Добавляем изображение
                parts.append(('image', img_id))
                
                # Обновляем позицию
                last_pos = end_pos
            
            # Добавляем оставшийся текст
            if last_pos < len(modified_html):
                parts.append(('text', modified_html[last_pos:]))
            
            # Обрабатываем каждую часть
            for part_type, content in parts:
                if part_type == 'text' and content.strip():
                    try:
                        # Создаем временный HTML без маркеров изображений
                        parser = HtmlToDocx()
                        
                        # Устанавливаем стиль для таблиц
                        # Проверяем наличие стиля TableGrid в документе
                        table_style = 'TableGrid'
                        if table_style not in docx_document.styles:
                            try:
                                # Пробуем создать стиль TableGrid
                                docx_document.styles.add_style('TableGrid', WD_STYLE_TYPE.TABLE)
                                logger.info("Создан стиль 'TableGrid' для таблиц")
                            except Exception:
                                # Если не удалось, используем Table Normal или оставляем без стиля
                                table_style = 'Table Normal' if 'Table Normal' in docx_document.styles else None
                                logger.warning(f"Не удалось создать стиль 'TableGrid', используем '{table_style or 'без стиля'}'")
                        
                        # Устанавливаем стиль таблицы
                        parser.table_style = table_style
                        
                        # Добавляем текст в документ
                        parser.add_html_to_document(content, docx_document)
                        
                        # Применяем отступы к созданным параграфам
                        for paragraph, p_id in parser.paragraphs_created:
                            # Проверяем, есть ли информация об отступах для этого параграфа
                            if p_id in indent_info:
                                info = indent_info[p_id]
                                
                                # Получаем текст параграфа
                                text = paragraph.text
                                
                                # Применяем отступ первой строки из text-indent
                                if info['text_indent']:
                                    value, unit = info['text_indent']
                                    try:
                                        # Преобразуем значение в сантиметры
                                        cm_value = convert_to_cm(float(value), unit)
                                        # Устанавливаем отступ первой строки напрямую
                                        paragraph.paragraph_format.first_line_indent = Cm(cm_value)
                                        logger.info(f"Установлен отступ первой строки {cm_value} см для параграфа {p_id}")
                                    except (ValueError, TypeError) as e:
                                        logger.warning(f"Не удалось применить отступ первой строки: {e}")
                                
                                # Применяем отступ слева из margin-left
                                if info['margin_left']:
                                    value, unit = info['margin_left']
                                    try:
                                        # Преобразуем значение в сантиметры
                                        cm_value = convert_to_cm(float(value), unit)
                                        # Устанавливаем отступ слева напрямую
                                        paragraph.paragraph_format.left_indent = Cm(cm_value)
                                        logger.info(f"Установлен отступ слева {cm_value} см для параграфа {p_id}")
                                    except (ValueError, TypeError) as e:
                                        logger.warning(f"Не удалось применить отступ слева: {e}")
                            
                            # Проверяем, есть ли информация о неразрывных пробелах для этого параграфа
                            auto_id = f"auto_{id(paragraph)}"
                            if p_id in paragraphs_with_nbsp or auto_id in paragraphs_with_nbsp:
                                nbsp_count = paragraphs_with_nbsp.get(p_id, paragraphs_with_nbsp.get(auto_id, 0))
                                if nbsp_count > 0:
                                    # Получаем текст параграфа без начальных неразрывных пробелов
                                    text = paragraph.text
                                    text_without_nbsp = text.lstrip('\u00A0')
                                    # Устанавливаем текст без неразрывных пробелов
                                    paragraph.text = text_without_nbsp
                                    # Устанавливаем отступ первой строки напрямую
                                    paragraph.paragraph_format.first_line_indent = Cm(nbsp_count * 0.25)
                                    logger.info(f"Установлен отступ первой строки {nbsp_count * 0.25} см на основе неразрывных пробелов")
                    except Exception as e:
                        logger.error(f"Ошибка при обработке текстовой части: {e}", exc_info=True)
                        # Добавляем текст напрямую в случае ошибки
                        clean_text = BeautifulSoup(content, 'html.parser').get_text()
                        if clean_text.strip():
                            docx_document.add_paragraph(clean_text)
                
                elif part_type == 'image':
                    img_id = content
                    img_path = images_map.get(img_id)
                    
                    if img_path and os.path.exists(img_path):
                        # Добавляем изображение в новый параграф
                        p = docx_document.add_paragraph()
                        run = p.add_run()
                        try:
                            # Устанавливаем выравнивание параграфа по центру для изображений
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Определяем размер изображения
                            try:
                                # Проверяем, доступен ли модуль PIL
                                try:
                                    from PIL import Image
                                    has_pil = True
                                except ImportError:
                                    has_pil = False
                                    logger.warning("Модуль PIL (Pillow) не установлен, размер изображения не будет определен")
                                
                                if has_pil:
                                    with Image.open(img_path) as img:
                                        width, height = img.size
                                        # Ограничиваем ширину изображения до 5.5 дюймов (14 см)
                                        max_width = Inches(5.5)
                                        if width > max_width.pt:
                                            # Сохраняем пропорции
                                            aspect_ratio = height / width
                                            width_inches = min(width / 72, 5.5)  # 72 DPI для преобразования в дюймы
                                            run.add_picture(img_path, width=Inches(width_inches))
                                        else:
                                            # Если изображение небольшое, добавляем его как есть
                                            run.add_picture(img_path)
                                else:
                                    # Если PIL не доступен, используем стандартную ширину
                                    run.add_picture(img_path, width=Inches(5.0))
                            except Exception as img_error:
                                # Если не удалось определить размер, используем стандартную ширину
                                logger.warning(f"Не удалось определить размер изображения: {img_error}")
                                run.add_picture(img_path, width=Inches(5.0))
                                
                            logger.info(f"Добавлено изображение: {img_path}")
                        except Exception as pic_error:
                            logger.error(f"Ошибка при добавлении изображения: {pic_error}")
                            # Пытаемся добавить изображение альтернативным способом
                            try:
                                run.add_picture(img_path)
                                logger.info(f"Изображение добавлено альтернативным способом: {img_path}")
                            except Exception as alt_error:
                                logger.error(f"Не удалось добавить изображение: {alt_error}")
                                # Добавляем текст с путем к изображению
                                run.add_text(f"[Изображение: {os.path.basename(img_path)}]")
            
            logger.info("HTML успешно преобразован и добавлен в документ")
            
            # ПРИНУДИТЕЛЬНО устанавливаем отступ для всех параграфов
            for paragraph in docx_document.paragraphs:
                try:
                    # Проверяем, что параграф не пустой
                    if paragraph.text.strip():
                        # Проверяем наличие стиля
                        style_name = getattr(paragraph.style, 'name', '') if hasattr(paragraph, 'style') and paragraph.style else ''
                        # Пропускаем заголовки
                        if style_name not in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title']:
                            paragraph.paragraph_format.first_line_indent = Cm(1.25)
                            logger.info(f"ПРИНУДИТЕЛЬНО установлен отступ 1.25 см для параграфа: '{paragraph.text[:20]}...'")
                except Exception as style_error:
                    logger.warning(f"Ошибка при установке отступа для параграфа: {style_error}")
    
    except Exception as e:
        logger.error(f"Ошибка при обработке HTML: {e}", exc_info=True)
        # В случае ошибки добавляем просто текст как запасной вариант
        try:
            plain_text = BeautifulSoup(html_content, 'html.parser').get_text()
            docx_document.add_paragraph(f"[Ошибка конвертации HTML, вставлен как текст]:\n{plain_text}")
        except Exception as fallback_e:
            logger.error(f"Ошибка при аварийной вставке HTML как текста: {fallback_e}")
            docx_document.add_paragraph("[Ошибка конвертации HTML, не удалось вставить даже как текст]")

def get_docx_template(template_name):
    """
    Возвращает путь к шаблону DOCX.
    
    Args:
        template_name (str): Название шаблона
        
    Returns:
        str: Путь к файлу шаблона или None, если шаблон не найден
    """
    template_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', f"{template_name}.docx")
    if os.path.exists(template_path):
        logger.info(f"Шаблон найден: {template_path}")
        return template_path
    
    logger.warning(f"Шаблон не найден: {template_path}")
    # Используем стандартный шаблон, если нужный не найден
    default_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', 'diplo_project.docx')
    if os.path.exists(default_path):
        logger.info(f"Используется стандартный шаблон: {default_path}")
        return default_path
    
    logger.error("Стандартный шаблон не найден")
    return None

def ensure_basic_styles(docx):
    """
    Проверяет наличие и создает базовые стили для документа.
    Применяет форматирование по ГОСТ: поля (левое - 30 мм, верхнее и нижнее - 20 мм, 
    правое - 10 мм), шрифт Times New Roman 12 пт, междустрочный интервал 1.5.
    
    Args:
        docx (Document): Документ python-docx
    """
    logger.info("Настройка базовых стилей и форматирования документа по ГОСТ")
    
    # Устанавливаем размеры полей для всего документа
    for section in docx.sections:
        section.left_margin = Cm(3.0)    # левое - 30 мм
        section.right_margin = Cm(1.0)   # правое - 10 мм
        section.top_margin = Cm(2.0)     # верхнее - 20 мм
        section.bottom_margin = Cm(2.0)  # нижнее - 20 мм
        logger.info("Установлены размеры полей: левое - 30 мм, правое - 10 мм, верхнее и нижнее - 20 мм")
    
    # Проверяем и создаем стиль Normal, если его нет
    if 'Normal' not in docx.styles:
        logger.info("Создание стиля 'Normal'")
        try:
            normal_style = docx.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
        except Exception as e:
            logger.warning(f"Не удалось создать стиль 'Normal', используем стандартный: {e}")
            normal_style = docx.styles['Normal']
    else:
        normal_style = docx.styles['Normal']
    
    # Устанавливаем параметры стиля Normal
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.first_line_indent = Cm(1.25)  # Отступ первой строки - 1.25 см
    normal_style.paragraph_format.line_spacing = 1.5           # Междустрочный интервал - 1.5
    normal_style.paragraph_format.space_after = Pt(0)          # Нет отступа после абзаца
    logger.info("Установлены параметры стиля Normal: Times New Roman, 12pt, интервал 1.5")
    
    # Создаем стили заголовков
    heading_styles = {
        'Heading 1': {'size': Pt(16), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
        'Heading 2': {'size': Pt(14), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.LEFT},
        'Heading 3': {'size': Pt(13), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.LEFT},
        'Heading 4': {'size': Pt(12), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.LEFT}
    }
    
    for style_name, params in heading_styles.items():
        try:
            if style_name not in docx.styles:
                style = docx.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            else:
                style = docx.styles[style_name]
                
            style.font.name = 'Times New Roman'
            style.font.size = params['size']
            style.font.bold = params['bold']
            style.paragraph_format.alignment = params['align']
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.5
            logger.info(f"Настроен стиль '{style_name}'")
        except Exception as e:
            logger.error(f"Ошибка при создании стиля '{style_name}': {e}")
    
    # Добавляем стили для списков
    list_styles = ['List Bullet', 'List Number']
    for list_style_name in list_styles:
        try:
            if list_style_name not in docx.styles:
                style = docx.styles.add_style(list_style_name, WD_STYLE_TYPE.PARAGRAPH)
            else:
                style = docx.styles[list_style_name]
                
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style.paragraph_format.left_indent = Pt(18)  # Отступ слева для списка
            style.paragraph_format.first_line_indent = Pt(-18)  # Отрицательный отступ первой строки (для маркера)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.5
            logger.info(f"Настроен стиль '{list_style_name}'")
        except Exception as e:
            logger.error(f"Ошибка при создании стиля '{list_style_name}': {e}")
    
    # Добавляем стиль TableGrid для таблиц
    try:
        if 'TableGrid' not in docx.styles:
            table_style = docx.styles.add_style('TableGrid', WD_STYLE_TYPE.TABLE)
            logger.info("Создан стиль 'TableGrid' для таблиц")
        
        if 'Table Normal' not in docx.styles:
            docx.styles.add_style('Table Normal', WD_STYLE_TYPE.TABLE)
            logger.info("Создан стиль 'Table Normal' для таблиц")
    except Exception as e:
        logger.error(f"Ошибка при создании стилей таблиц: {e}")
    
    logger.info("Базовые стили документа настроены успешно")

def replace_document_fields(docx, replacements):
    """
    Заменяет поля в шаблоне документа их значениями с сохранением форматирования.
    При замене поля на значение добавляет нужное количество пробелов для сохранения
    размера строки, если значение короче плейсхолдера.
    
    Args:
        docx (Document): Документ python-docx
        replacements (dict): Словарь с заменами {поле: значение}
    """
    logger.info("Замена полей в шаблоне с сохранением форматирования")
    
    # Заменяем поля в параграфах более умным способом
    for paragraph in docx.paragraphs:
        # Проверяем, содержит ли параграф какое-либо поле
        for field, value in replacements.items():
            if field in paragraph.text:
                # Преобразуем значение поля в строку и обрабатываем специальные случаи
                value_str = str(value) if value is not None else ""
                
                # Вычисляем разницу в длине между полем и значением
                length_diff = len(field) - len(value_str)
                
                # Если значение короче поля, добавляем пробелы для компенсации
                if length_diff > 0:
                    # Добавляем пробелы, чтобы сохранить форматирование
                    value_str = value_str + ' ' * length_diff
                
                # Заменяем поле на значение с компенсированной длиной
                paragraph.text = paragraph.text.replace(field, value_str)
    
    # Заменяем поля в таблицах с тем же подходом
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for field, value in replacements.items():
                        if field in paragraph.text:
                            # Преобразуем значение поля в строку
                            value_str = str(value) if value is not None else ""
                            
                            # Вычисляем разницу в длине
                            length_diff = len(field) - len(value_str)
                            
                            # Если значение короче поля, добавляем пробелы
                            if length_diff > 0:
                                value_str = value_str + ' ' * length_diff
                            
                            # Заменяем поле на компенсированное значение
                            paragraph.text = paragraph.text.replace(field, value_str)

def apply_document_formatting(docx, standard_name):
    """
    Применяет стили форматирования к документу по ГОСТ.
    Эта функция теперь просто вызывает ensure_basic_styles,
    так как все форматирование уже настроено там.
    
    Args:
        docx (Document): Документ python-docx
        standard_name (str): Название стандарта (например, "ГОСТ 7.32-2017")
    """
    logger.info(f"Применение форматирования по стандарту: {standard_name}")
    
    # Просто вызываем ensure_basic_styles, так как все форматирование уже настроено там
    ensure_basic_styles(docx)
    
    # Дополнительно можно добавить специфичные настройки для конкретных стандартов
    if "7.32" in standard_name:  # ГОСТ 7.32
        logger.info("Применены специфические настройки для ГОСТ 7.32")
        # Здесь можно добавить специфичные настройки для ГОСТ 7.32, если нужно
    elif "2.105" in standard_name:  # ГОСТ 2.105
        logger.info("Применены специфические настройки для ГОСТ 2.105")
        # Здесь можно добавить специфичные настройки для ГОСТ 2.105, если нужно
    
    logger.info("Форматирование документа успешно применено")

def add_page_numbers(docx):
    """
    Добавляет нумерацию страниц в документ в нижней части страницы по центру.
    
    Args:
        docx (Document): Документ python-docx
    """
    logger.info("Добавление нумерации страниц")
    
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        section = docx.sections[0]
        footer = section.footer
        
        # Очищаем существующее содержимое колонтитула
        for p in footer.paragraphs:
            p._element.getparent().remove(p._element)
            p._p = None
            p._element = None
        
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Функция для создания номера страницы в Word
        def create_element(name):
            return OxmlElement(name)
        
        def create_attribute(element, name, value):
            element.set(qn(name), value)
        
        # Добавляем номер страницы
        run = paragraph.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        
        # Создаем поле PAGE
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
        
        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')
        
        # Получаем элемент для запуска 
        r_element = run._r
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        
        logger.info("Нумерация страниц успешно добавлена")
    except Exception as e:
        logger.error(f"Ошибка при добавлении нумерации страниц: {e}")

def merge_docs(doc_title, doc_body):
    """
    Объединяет два документа: копирует все элементы из doc_body в конец doc_title.
    Обеспечивает корректное копирование всех элементов, включая изображения.
    
    Args:
        doc_title (Document): Документ-приемник (с титульным листом)
        doc_body (Document): Документ-источник (с основным содержимым)
    """
    import copy
    from docx.oxml.section import CT_SectPr
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    
    try:
        # Получаем тело документа-приемника
        body_target = doc_title.element.body
        
        # Копируем все стили из документа-источника в документ-приемник
        # Объект Styles не поддерживает метод items(), поэтому обходим стили по-другому
        try:
            # Получаем все стили из документа-источника
            source_styles = doc_body.styles
            # Получаем все стили из документа-приемника
            target_styles = doc_title.styles
            
            # Проходим по всем стилям в документе-источнике
            for style in source_styles:
                style_id = style.style_id
                # Проверяем, существует ли стиль в документе-приемнике
                if style_id not in [s.style_id for s in target_styles]:
                    try:
                        # Добавляем стиль в документ-приемник
                        target_styles.add_style(style_id, style.type, style.name)
                        logger.info(f"Скопирован стиль {style_id} из документа-источника")
                    except Exception as style_error:
                        logger.warning(f"Не удалось скопировать стиль {style_id}: {style_error}")
        except Exception as styles_error:
            logger.warning(f"Ошибка при копировании стилей: {styles_error}")
            # Продолжаем выполнение даже если не удалось скопировать стили
        
        # Создаем словарь для отображения старых ID изображений на новые
        image_map = {}
        
        # Копируем все изображения из документа-источника
        # Для этого нужно скопировать все отношения (relationships)
            # Копируем все изображения из документа-источника
        for rel_id, rel in list(doc_body.part.rels.items()):
            if rel.is_external:
                logger.info(f"Пропущено внешнее отношение: {rel_id}")
                continue

            if rel.reltype == RT.IMAGE:
                try:
                    # Получаем исходную часть изображения
                    source_image_part = rel.target_part
                    image_bytes = source_image_part.blob

                    # Добавляем (или находим) эту картинку в пакете целевого документа
                    new_image_part = doc_title.part.package.image_parts.get_or_add_image_part(
                        io.BytesIO(image_bytes)
                    )

                    # Создаём отношение в основном part'е и сразу получаем rId
                    new_rId = doc_title.part.relate_to(
                        new_image_part,
                        RT.IMAGE
                    )

                    image_map[rel_id] = new_rId
                    logger.info(
                        f"Скопировано/использовано изображение {rel_id} (старый) -> "
                        f"{new_rId} (новый), часть: {new_image_part.partname}, "
                        f"тип: {new_image_part.content_type}"
                    )
                except Exception as img_error:
                    logger.warning(
                        f"Ошибка при копировании изображения {rel_id}: {img_error}",
                        exc_info=True
                    )
        
        # Копируем все элементы из документа-источника
        for element in doc_body.element.body:
            # Пропускаем секционные свойства, чтобы не нарушить структуру документа
            if isinstance(element, CT_SectPr):
                logger.info("Пропущены секционные свойства при объединении документов")
                continue
                
            # Создаем глубокую копию XML-узла
            new_el = copy.deepcopy(element)
            
            # Обновляем ссылки на изображения в новом элементе
            try:
                # Находим все элементы blip в XML-дереве без использования xpath с namespaces
                # Используем прямой обход XML-дерева
                def update_blip_refs(element):
                    # Проверяем, является ли элемент blip
                    if element.tag.endswith('}blip'):
                        # Получаем атрибут с namespace для embed
                        for key, value in element.attrib.items():
                            if key.endswith('}embed') and value in image_map:
                                # Заменяем на новый ID
                                element.attrib[key] = image_map[value]
                                logger.info(f"Обновлена ссылка на изображение: {value} -> {image_map[value]}")
                    
                    # Рекурсивно обходим все дочерние элементы
                    for child in element:
                        update_blip_refs(child)
                
                # Запускаем обход с корня элемента
                update_blip_refs(new_el)
            except Exception as blip_error:
                logger.warning(f"Ошибка при обновлении ссылок на изображения: {blip_error}")
            
            # Добавляем элемент в документ-приемник
            body_target.append(new_el)
        
        logger.info("Документы успешно объединены")
    except Exception as e:
        logger.error(f"Ошибка при объединении документов: {e}", exc_info=True)
        raise

def apply_formatting_to_paragraphs(doc):
    """
    Принудительно применяет форматирование ко всем параграфам документа.
    Эта функция вызывается после слияния документов, чтобы убедиться, 
    что все стили применены правильно.
    
    Args:
        doc (Document): Документ python-docx
    """
    logger.info("Принудительное применение форматирования ко всем параграфам")
    
    # Применяем форматирование ко всем параграфам
    for paragraph in doc.paragraphs:
        # Пропускаем пустые параграфы
        if not paragraph.text.strip():
            continue
            
        # Определяем, является ли параграф заголовком
        style_name = paragraph.style.name if paragraph.style else ''
        
        if style_name == 'Heading1':
            # Применяем форматирование заголовков Heading 1
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Применяем форматирование шрифта к каждому Run в параграфе
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.bold = True
                run.font.size = Pt(16)
        elif style_name.startswith('Heading'):
            # Применяем форматирование других заголовков
            paragraph.paragraph_format.line_spacing = 1.5
            
            # Применяем форматирование шрифта к каждому Run в параграфе
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.bold = True
                
                # Размер в зависимости от уровня заголовка
                if style_name == 'Heading2':
                    run.font.size = Pt(14)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif style_name == 'Heading3':
                    run.font.size = Pt(13)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    run.font.size = Pt(12)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Применяем форматирование обычного текста
            try:
                # Устанавливаем шрифт и размер для всех элементов текста
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                
                # Устанавливаем параметры абзаца
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Cm(1.25)
                paragraph.paragraph_format.line_spacing = 1.5
                paragraph.paragraph_format.space_after = Pt(0)
            except Exception as e:
                logger.warning(f"Не удалось применить форматирование к параграфу: {e}")
    
    logger.info("Форматирование применено ко всем параграфам документа")

@login_required
def main_export_docx(request, pk):
    """
    Экспорт документа Main в формат DOCX.
    
    Args:
        request: HTTP запрос
        pk (int): ID документа
        
    Returns:
        HttpResponse: Ответ с DOCX файлом
    """
    logger.info(f"Начат экспорт документа Main (ID: {pk}) в DOCX")
    
    # Проверяем наличие необходимых зависимостей
    try:
        from PIL import Image
        logger.info("Модуль PIL (Pillow) доступен")
    except ImportError:
        logger.warning("Модуль PIL (Pillow) не установлен. Рекомендуется установить для корректной обработки изображений")
        # Проверяем наличие файла requirements.txt
        requirements_path = os.path.join(settings.BASE_DIR, 'requirements.txt')
        if os.path.exists(requirements_path):
            logger.info("Добавление Pillow в requirements.txt")
            try:
                with open(requirements_path, 'r') as f:
                    requirements = f.read()
                
                # Проверяем, есть ли уже Pillow в requirements.txt
                if 'Pillow' not in requirements and 'pillow' not in requirements:
                    with open(requirements_path, 'a') as f:
                        f.write("\n# Добавлено автоматически для обработки изображений\nPillow>=9.0.0\n")
                    logger.info("Pillow добавлен в requirements.txt")
                    messages.info(request, "Для корректной обработки изображений рекомендуется установить Pillow: pip install Pillow")
            except Exception as req_error:
                logger.warning(f"Не удалось обновить requirements.txt: {req_error}")
    
    try:
        # Получаем документ
        document = get_object_or_404(Document_main, pk=pk, owner=request.user)
        logger.info(f"Документ найден: {document.title}")
        
        # Определяем шаблон на основе типа работы
        template_name = WORK_TYPE_TEMPLATES.get(document.work_type, 'diplo_project')
        template_path = get_docx_template(template_name)
        
        if not template_path:
            messages.error(request, "Не удалось найти шаблон для документа.")
            return redirect('documents:main_detail', pk=pk)
        
        # Шаг 1: Загружаем титульный документ и рендерим его с помощью DocxTemplate
        doc_template = DocxTemplate(template_path)
        
        # Подготавливаем контекст для шаблона
        context = {
            'TITLE': document.title.upper(),
            'TitleContinue': "",
            'YEAR': str(document.year),
            'YearShort': str(document.year)[-2:] if document.year else '__',
            'STUDENT_NAME': document.student_name or '',
            'SUPERVISOR': document.supervisor or '',
            'SupervisorPosition': getattr(document, 'supervisor_position', ''),
            'SupervisorSignature': '_________',
            'StudentSignature': '_________',
            'Institut': document.institute_name or 'Институт космических и информационных технологий',
            'institut': document.institute_name or 'Институт космических и информационных технологий',  # вариант с маленькой буквы
            'Kafedra': document.department_name or '',
            'ZavKaf': getattr(document, 'head_of_department', ''),
            'Podpis': '_________',
            'Day': getattr(document, 'day', '___'),
            'Month': getattr(document, 'month', '________'),
            'Speciality': f"{document.specialty_code} {document.specialty_name}".strip(),
            'UNIVERSITY': (document.university_name or 'СИБИРСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ').upper(),
            'code': document.specialty_code,
            'head_of_department': getattr(document, 'head_of_department', ''),
            'speciality_full': f"{document.specialty_code_full} {document.specialty_name}".strip(),
            'record_number': document.record_number,
            'reviewer': document.reviewer or '',
            'reviewer_position': getattr(document, 'reviewer_position', ''),
            'factory_supervisor': document.factory_supervisor or '',
        }
        
        # Рендерим документ с указанным контекстом
        logger.info("Заполнение шаблона через DocxTemplate")
        doc_template.render(context)
        
        # Сохраняем титульный лист в BytesIO вместо временного файла
        title_io = BytesIO()
        doc_template.save(title_io)
        title_io.seek(0)
        logger.info("Титульный лист сохранен в BytesIO")
        
        # Загружаем сохраненный титульный лист как Document
        doc_title = Document(title_io)
        logger.info("Титульный лист успешно подготовлен")
        
        # Шаг 2: Создаем отдельный документ для основной части
        doc_body = Document()
        
        # Проверяем и создаем базовые стили только для основного текста
        ensure_basic_styles(doc_body)
        
        # Применяем стили форматирования на основе стандарта (только для основного текста)
        if document.standart:
            logger.info(f"Применение стилей форматирования для стандарта: {document.standart}")
            apply_document_formatting(doc_body, document.standart)
        
        # Добавляем содержимое документа
        if document.data:
            logger.info("Добавление содержимого документа")
            process_html_to_docx(document.data, doc_body)
        else:
            doc_body.add_paragraph("Документ не содержит данных")
        
        # Добавляем список литературы
        add_references_section(doc_body, document)
            
        # Добавляем нумерацию страниц
        add_page_numbers(doc_body)
        
        # Сохраняем основную часть в BytesIO вместо временного файла
        body_io = BytesIO()
        doc_body.save(body_io)
        body_io.seek(0)
        logger.info("Основная часть сохранена в BytesIO")
        
        # Шаг 3: Объединяем документы
        logger.info("Объединение титульного листа и основной части")
        merge_docs(doc_title, Document(body_io))
        
        # Шаг 4: Применяем форматирование после слияния
        logger.info("Применение форматирования после слияния документов")
        # Устанавливаем поля для всех секций
        for section in doc_title.sections:
            section.left_margin = Cm(3.0)    # левое - 30 мм
            section.right_margin = Cm(1.0)   # правое - 10 мм
            section.top_margin = Cm(2.0)     # верхнее - 20 мм
            section.bottom_margin = Cm(2.0)  # нижнее - 20 мм
        
        # Применяем форматирование ко всем параграфам
        apply_formatting_to_paragraphs(doc_title)
        
        # Шаг 5: Сохраняем итоговый документ в BytesIO
        final_io = BytesIO()
        doc_title.save(final_io)
        final_io.seek(0)
        logger.info("Итоговый документ сохранен в BytesIO")
        
        # Проверяем, что данные есть
        file_content = final_io.getvalue()
        if not file_content or len(file_content) == 0:
            logger.error("Итоговый файл имеет нулевой размер")
            messages.error(request, "Ошибка при создании документа: файл не был создан")
            return redirect('documents:main_detail', pk=pk)
        
        # Формируем имя файла
        safe_filename = slugify(document.document_name or document.title)
        if not safe_filename:  # Дополнительная проверка на пустое имя
            safe_filename = f"document_{pk}"
        filename = f"{safe_filename}.docx"
        
        # Отправляем файл пользователю с правильным Content-Disposition
        response = HttpResponse(
            file_content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Экспорт документа Main в DOCX успешно завершен. Имя файла: {filename}")
        return response
        
    except Exception as e:
        logger.error(f"Ошибка при экспорте документа Main в DOCX: {e}", exc_info=True)
        messages.error(request, f"Ошибка при экспорте документа: {str(e)}")
        return redirect('documents:main_detail', pk=pk)

@login_required
def main_export_pdf(request, pk):
    """
    Экспорт документа Main в формат PDF.
    
    Args:
        request: HTTP запрос
        pk (int): ID документа
        
    Returns:
        HttpResponse: Ответ с PDF файлом
    """
    logger.info(f"Начат экспорт документа Main (ID: {pk}) в PDF") 
    
    # Проверяем наличие необходимых зависимостей
    try:
        from PIL import Image
        logger.info("Модуль PIL (Pillow) доступен")
    except ImportError:
        logger.warning("Модуль PIL (Pillow) не установлен. Рекомендуется установить для корректной обработки изображений")
        # Проверяем наличие файла requirements.txt
        requirements_path = os.path.join(settings.BASE_DIR, 'requirements.txt')
        if os.path.exists(requirements_path):
            logger.info("Добавление Pillow в requirements.txt")
            try:
                with open(requirements_path, 'r') as f:
                    requirements = f.read()
                
                # Проверяем, есть ли уже Pillow в requirements.txt
                if 'Pillow' not in requirements and 'pillow' not in requirements:
                    with open(requirements_path, 'a') as f:
                        f.write("\n# Добавлено автоматически для обработки изображений\nPillow>=9.0.0\n")
                    logger.info("Pillow добавлен в requirements.txt")
                    messages.info(request, "Для корректной обработки изображений рекомендуется установить Pillow: pip install Pillow")
            except Exception as req_error:
                logger.warning(f"Не удалось обновить requirements.txt: {req_error}")
    
    com_initialized = False
    if pythoncom:
        try:
            pythoncom.CoInitialize()
            com_initialized = True
            logger.info("COM успешно инициализирован.")
        except Exception as e:
            logger.warning(f"Ошибка при инициализации COM: {e}")
            pass
            
    try:
        document_obj = get_object_or_404(Document_main, pk=pk, owner=request.user)
        logger.info(f"Документ найден: {document_obj.title}")
        
        # Для PDF нам нужны реальные файлы для конвертации docx2pdf
        # Но мы будем использовать BytesIO где возможно и минимизировать использование временных файлов
        
        # Создаем DOCX документ в памяти
        title_io = BytesIO()
        body_io = BytesIO()
        
        template_name = WORK_TYPE_TEMPLATES.get(document_obj.work_type, 'diplo_project')
        template_path = get_docx_template(template_name)
        
        if not template_path:
            messages.error(request, "Не удалось найти шаблон для документа.")
            return redirect('documents:main_detail', pk=pk)
        
        # 1. Рендеринг шаблона DocxTemplate для титульного листа
        doc_template = DocxTemplate(template_path)
        context = {
            'TITLE': document_obj.title.upper(),
            'TitleContinue': "",
            'YEAR': str(document_obj.year),
            'YearShort': str(document_obj.year)[-2:] if document_obj.year else '__',
            'STUDENT_NAME': document_obj.student_name or '',
            'SUPERVISOR': document_obj.supervisor or '',
            'SupervisorPosition': getattr(document_obj, 'supervisor_position', ''),
            'SupervisorSignature': '_________',
            'StudentSignature': '_________',
            'Institut': document_obj.institute_name or 'Институт космических и информационных технологий',
            'institut': document_obj.institute_name or 'Институт космических и информационных технологий',
            'Kafedra': document_obj.department_name or '',
            'ZavKaf': getattr(document_obj, 'head_of_department', ''),
            'Podpis': '_________',
            'Day': getattr(document_obj, 'day', '___'),
            'Month': getattr(document_obj, 'month', '________'),
            'Speciality': f"{document_obj.specialty_code} {document_obj.specialty_name}".strip(),
            'UNIVERSITY': (document_obj.university_name or 'СИБИРСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ').upper(),
            'code': document_obj.specialty_code,
            'head_of_department': getattr(document_obj, 'head_of_department', ''),
            'speciality_full': f"{document_obj.specialty_code_full} {document_obj.specialty_name}".strip(),
            'record_number': document_obj.record_number,
            'reviewer': document_obj.reviewer or '',
            'reviewer_position': getattr(document_obj, 'reviewer_position', ''),
            'factory_supervisor': document_obj.factory_supervisor or '',
        }
        logger.info("Заполнение шаблона через DocxTemplate")
        doc_template.render(context)
        
        # Сохраняем титульный лист в BytesIO
        doc_template.save(title_io)
        title_io.seek(0)
        logger.info("Титульный лист сохранен в BytesIO")
        
        # Загружаем сохраненный титульный лист как Document
        doc_title = Document(title_io)
        logger.info("Титульный лист успешно подготовлен")
        
        # 2. Создаем отдельный документ для основной части
        doc_body = Document()
        
        # Проверяем и создаем базовые стили для основного текста
        ensure_basic_styles(doc_body)
        
        # Применяем стили форматирования на основе стандарта
        if document_obj.standart:
            logger.info(f"Применение стилей форматирования для стандарта: {document_obj.standart}")
            apply_document_formatting(doc_body, document_obj.standart)
        
        # Добавляем содержимое документа
        if document_obj.data:
            logger.info("Добавление содержимого документа")
            process_html_to_docx(document_obj.data, doc_body)
        else:
            doc_body.add_paragraph("Документ не содержит данных")
        
        # Добавляем список литературы
        add_references_section(doc_body, document_obj)
            
        # Добавляем нумерацию страниц
        add_page_numbers(doc_body)
        
        # Сохраняем основную часть в BytesIO
        doc_body.save(body_io)
        body_io.seek(0)
        logger.info("Основная часть сохранена в BytesIO")
        
        # 3. Объединяем документы
        logger.info("Объединение титульного листа и основной части")
        merge_docs(doc_title, Document(body_io))
        
        # 4. Применяем форматирование после слияния
        logger.info("Применение форматирования после слияния документов")
        # Устанавливаем поля для всех секций
        for section in doc_title.sections:
            section.left_margin = Cm(3.0)    # левое - 30 мм
            section.right_margin = Cm(1.0)   # правое - 10 мм
            section.top_margin = Cm(2.0)     # верхнее - 20 мм
            section.bottom_margin = Cm(2.0)  # нижнее - 20 мм
        
        # Применяем форматирование ко всем параграфам
        apply_formatting_to_paragraphs(doc_title)
        
        # Для конвертации в PDF нам нужен временный файл DOCX
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_temp:
            docx_path = docx_temp.name
            logger.info(f"Сохранение объединенного документа во временный файл: {docx_path}")
            
            # Сохраняем объединенный документ во временный файл
            doc_title.save(docx_path)
            
            # Закрываем файл перед конвертацией
            docx_temp.close()
            
            # Проверяем, что файл создан и имеет размер
            if not os.path.exists(docx_path) or os.path.getsize(docx_path) == 0:
                logger.error("Объединенный DOCX файл не создан или имеет нулевой размер")
                messages.error(request, "Ошибка при создании документа: файл не был создан")
                # Удаляем временный файл
                try:
                    os.unlink(docx_path)
                except:
                    pass
                return redirect('documents:main_detail', pk=pk)
            
            # Создаем временный файл для PDF
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
                pdf_path = pdf_temp.name
                pdf_temp.close()  # Закрываем файл перед конвертацией
                
                # Конвертация в PDF
                logger.info(f"Конвертация DOCX в PDF: {docx_path} -> {pdf_path}")
                try:
                    convert(docx_path, pdf_path)
                    logger.info(f"PDF файл создан: {pdf_path}")
                    
                    # Проверяем, что PDF файл создан и имеет размер
                    if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
                        logger.error("PDF файл не создан или имеет нулевой размер")
                        messages.error(request, "Ошибка при создании PDF: файл не был создан")
                        # Удаляем временные файлы
                        try:
                            os.unlink(docx_path)
                            os.unlink(pdf_path)
                        except:
                            pass
                        return redirect('documents:main_detail', pk=pk)
                    
                    safe_filename = slugify(document_obj.document_name or document_obj.title or "document")
                    if not safe_filename:
                        safe_filename = f"document_{pk}"
                    filename = f"{safe_filename}.pdf" 
                    
                    # Читаем PDF в память и отправляем пользователю
                    with open(pdf_path, 'rb') as pdf_file:
                        pdf_content = pdf_file.read()
                        
                    # Удаляем временные файлы
                    try:
                        os.unlink(docx_path)
                        os.unlink(pdf_path)
                    except Exception as cleanup_error:
                        logger.warning(f"Не удалось удалить временные файлы: {cleanup_error}")
                    
                    # Отправляем PDF пользователю
                    response = HttpResponse(pdf_content, content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    
                    logger.info(f"Экспорт документа Main в PDF успешно завершен. Имя файла: {filename}")
                    return response
                    
                except Exception as pdf_error:
                    logger.error(f"Ошибка при конвертации в PDF: {pdf_error}", exc_info=True)
                    messages.error(request, f"Ошибка при конвертации в PDF: {str(pdf_error)}")
                    
                    # Удаляем временные файлы
                    try:
                        os.unlink(docx_path)
                        os.unlink(pdf_path)
                    except:
                        pass
                    
                    # В случае ошибки конвертации предлагаем скачать DOCX
                    safe_filename = slugify(document_obj.document_name or document_obj.title or "document")
                    if not safe_filename:
                        safe_filename = f"document_{pk}"
                    filename = f"{safe_filename}.docx"
                    
                    # Читаем DOCX в память
                    with open(docx_path, 'rb') as docx_file:
                        docx_content = docx_file.read()
                    
                    # Отправляем DOCX пользователю
                    response = HttpResponse(
                        docx_content,
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    
                    logger.info(f"Предоставлен DOCX файл вместо PDF из-за ошибки конвертации. Имя файла: {filename}")
                    return response
                
    except Exception as e:
        logger.error(f"Ошибка при экспорте документа Main в PDF: {e}", exc_info=True)
        messages.error(request, f"Ошибка при экспорте документа: {str(e)}")
        return redirect('documents:main_detail', pk=pk)
    finally:
        if com_initialized:
            try:
                pythoncom.CoUninitialize()
                logger.info("COM успешно деинициализирован.")
            except Exception as e:
                logger.warning(f"Ошибка при деинициализации COM: {e}")

def get_metadata_from_doi(doi):
    """
    Получает метаданные публикации по DOI через CrossRef API.
    
    Args:
        doi (str): DOI публикации
        
    Returns:
        dict: Словарь с метаданными или None в случае ошибки
    """
    logger.info(f"Получение метаданных для DOI: {doi}")
    
    try:
        # Формируем URL для запроса к CrossRef API
        url = f"https://api.crossref.org/works/{doi}"
        headers = {
            "Accept": "application/json",
            "User-Agent": "GostDocsApp/1.0 (mailto:admin@example.com)"
        }
        
        response = requests.get(url, headers=headers)
        
        if response.status_code == 200:
            data = response.json()
            if 'message' in data:
                logger.info(f"Метаданные для DOI {doi} успешно получены")
                return data['message']
            else:
                logger.warning(f"В ответе CrossRef API отсутствует поле 'message' для DOI {doi}")
                return None
        else:
            logger.warning(f"Ошибка при запросе к CrossRef API для DOI {doi}: {response.status_code}")
            return None
            
    except Exception as e:
        logger.error(f"Ошибка при получении метаданных для DOI {doi}: {e}", exc_info=True)
        return None

def format_citation_gost(metadata):
    """
    Форматирует библиографическую ссылку по ГОСТ 7.1-2003 на основе метаданных CrossRef.
    
    Args:
        metadata (dict): Метаданные публикации из CrossRef API
        
    Returns:
        str: Отформатированная библиографическая ссылка
    """
    try:
        # Получаем основные данные из метаданных
        title = metadata.get('title', [''])[0] if isinstance(metadata.get('title', []), list) else metadata.get('title', '')
        
        # Получаем авторов
        authors = []
        if 'author' in metadata:
            for author in metadata['author']:
                given = author.get('given', '')
                family = author.get('family', '')
                if given and family:
                    # Формат: Фамилия И.О.
                    initials = ''.join([name[0] + '.' for name in given.split()])
                    authors.append(f"{family} {initials}")
                elif family:
                    authors.append(family)
        
        # Формируем строку авторов по ГОСТ
        authors_str = ""
        if authors:
            if len(authors) == 1:
                authors_str = authors[0]
            elif len(authors) <= 3:
                authors_str = ", ".join(authors)
            else:
                authors_str = f"{authors[0]} и др."
        
        # Получаем данные о журнале или издательстве
        container_title = metadata.get('container-title', [''])[0] if isinstance(metadata.get('container-title', []), list) else metadata.get('container-title', '')
        publisher = metadata.get('publisher', '')
        
        # Получаем данные о выпуске
        volume = metadata.get('volume', '')
        issue = metadata.get('issue', '')
        page = metadata.get('page', '')
        
        # Получаем год публикации
        year = ""
        if 'published' in metadata and 'date-parts' in metadata['published']:
            year = str(metadata['published']['date-parts'][0][0]) if metadata['published']['date-parts'][0] else ""
        
        # Формируем ссылку по ГОСТ 7.1-2003
        citation = ""
        
        # Добавляем авторов
        if authors_str:
            citation += f"{authors_str}. "
        
        # Добавляем название
        if title:
            citation += f"{title}"
            # Добавляем точку, если название не заканчивается знаком препинания
            if not title[-1] in ['.', '!', '?']:
                citation += ". "
            else:
                citation += " "
        
        # Добавляем данные о журнале/издательстве
        if container_title:
            citation += f"// {container_title}. "
        elif publisher:
            citation += f"// {publisher}. "
        
        # Добавляем год
        if year:
            citation += f"{year}. "
        
        # Добавляем данные о выпуске
        if volume:
            citation += f"Т. {volume}"
            if issue or page:
                citation += ", "
        
        if issue:
            citation += f"№ {issue}"
            if page:
                citation += ", "
        
        if page:
            citation += f"С. {page}"
        
        # Добавляем DOI
        if 'DOI' in metadata:
            citation += f". DOI: {metadata['DOI']}"
        
        return citation
        
    except Exception as e:
        logger.error(f"Ошибка при форматировании ссылки по ГОСТ: {e}", exc_info=True)
        return f"Ошибка форматирования ссылки: {str(e)}"

def add_references_section(docx_document, document_obj):
    """
    Добавляет раздел со списком литературы в документ.
    
    Args:
        docx_document (Document): Документ python-docx
        document_obj: Объект документа из базы данных
    """
    logger.info("Добавление раздела со списком литературы")
    
    try:
        # Добавляем заголовок раздела
        docx_document.add_paragraph("", style='Normal')  # Пустая строка перед разделом
        heading = docx_document.add_paragraph("Список литературы", style='Heading 1')
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Получаем список DOI из документа
        # Предполагаем, что DOI хранятся в поле references_doi как строка с разделителями
        doi_list = []
        if hasattr(document_obj, 'references_doi') and document_obj.references_doi:
            # Разбиваем строку на отдельные DOI
            doi_list = [doi.strip() for doi in document_obj.references_doi.split(',') if doi.strip()]
        
        # Если список DOI пуст, добавляем информационное сообщение
        if not doi_list:
            docx_document.add_paragraph("Список литературы не содержит источников.", style='Normal')
            return
        
        # Обрабатываем каждый DOI и добавляем ссылку в документ
        for i, doi in enumerate(doi_list, 1):
            # Получаем метаданные по DOI
            metadata = get_metadata_from_doi(doi)
            
            if metadata:
                # Форматируем ссылку по ГОСТ
                citation = format_citation_gost(metadata)
                # Добавляем ссылку в документ с номером
                p = docx_document.add_paragraph(style='Normal')
                p.add_run(f"{i}. ").bold = True
                p.add_run(citation)
            else:
                # Если не удалось получить метаданные, добавляем только DOI
                p = docx_document.add_paragraph(style='Normal')
                p.add_run(f"{i}. ").bold = True
                p.add_run(f"DOI: {doi} (не удалось получить метаданные)")
        
        logger.info(f"Добавлено {len(doi_list)} источников в список литературы")
        
    except Exception as e:
        logger.error(f"Ошибка при добавлении раздела со списком литературы: {e}", exc_info=True)
        docx_document.add_paragraph("Ошибка при формировании списка литературы", style='Normal')

def convert_to_cm(value, unit):
    """
    Преобразует значение из различных единиц измерения в сантиметры.
    
    Args:
        value (float): Числовое значение
        unit (str): Единица измерения (px, em, cm, mm, pt)
        
    Returns:
        float: Значение в сантиметрах
    """
    if unit == 'cm':
        return value
    elif unit == 'mm':
        return value / 10.0
    elif unit == 'pt':
        return value * 0.0352778  # 1 pt = 0.0352778 см
    elif unit == 'px':
        return value * 0.0264583  # Примерно 96 px = 2.54 см
    elif unit == 'em':
        return value * 0.42333  # Примерно 1 em = 12pt = 0.42333 см
    else:
        # По умолчанию предполагаем пиксели
        return value * 0.0264583
